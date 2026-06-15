import docx
import openpyxl
import os
import sys

# Ensure UTF-8 output even if printing
sys.stdout.reconfigure(encoding='utf-8')

def parse_docx(file_path, out_path):
    print(f"Parsing DOCX: {file_path}")
    doc = docx.Document(file_path)
    lines = []
    
    # Read paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Let's see if there is heading style
            if para.style.name.startswith('Heading'):
                level = para.style.name.replace('Heading', '').strip()
                if level.isdigit():
                    prefix = '#' * int(level)
                    lines.append(f"\n{prefix} {text}\n")
                else:
                    lines.append(f"\n# {text}\n")
            else:
                lines.append(text)
        else:
            lines.append("") # empty line
            
    # Read tables
    for i, table in enumerate(doc.tables):
        lines.append(f"\n### Table {i+1}\n")
        for r_idx, row in enumerate(table.rows):
            row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            # remove duplicates for merged cells if necessary, but simple join is fine for now
            row_str = " | ".join(row_data)
            lines.append(f"| {row_str} |")
            if r_idx == 0:
                # separator
                sep = " | ".join(["---"] * len(row.cells))
                lines.append(f"| {sep} |")
        lines.append("")
        
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"DOCX written to {out_path}")

def parse_xlsx(file_path, out_path):
    print(f"Parsing XLSX: {file_path}")
    wb = openpyxl.load_workbook(file_path, data_only=True)
    lines = []
    for sheet_name in wb.sheetnames:
        lines.append(f"\n# Sheet: {sheet_name}\n")
        sheet = wb[sheet_name]
        
        # We need to capture columns nicely
        for r_idx, row in enumerate(sheet.iter_rows(values_only=True)):
            if any(row is not None and str(x).strip() != "" for x in row if x is not None):
                # Clean row values
                row_vals = []
                for val in row:
                    if val is None:
                        row_vals.append("")
                    else:
                        # Format dates nicely
                        if hasattr(val, 'strftime'):
                            row_vals.append(val.strftime('%Y-%m-%d'))
                        else:
                            row_vals.append(str(val).strip().replace('\n', ' '))
                lines.append(" | ".join(row_vals))
                
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"XLSX written to {out_path}")

if __name__ == "__main__":
    parse_docx("docs/钢铁板材表面缺陷智能检测系统需求说明书_V2完善版.docx", "docs/spec_doc.md")
    parse_xlsx("docs/项目计划.xlsx", "docs/project_plan.md")
